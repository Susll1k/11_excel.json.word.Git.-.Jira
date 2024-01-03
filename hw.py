# -------------------------------1--------------------------------


import openpyxl

wb1 = openpyxl.load_workbook('file1.xlsx')
wb2 = openpyxl.load_workbook('file2.xlsx')
wb3 = openpyxl.load_workbook('file3.xlsx')

ws1 = wb1.active
ws2 = wb2.active
ws3 = wb3.active


ws1_sheet=ws1['A1'].value
ws2_sheet=ws2['A1'].value
ws3_sheet=ws3['A1'].value
lst=sorted([ws1_sheet, ws2_sheet, ws3_sheet], reverse=True)

workbook = openpyxl.Workbook()
sheet = workbook.active


sheet['A1'] = lst[0]
sheet['A2'] = lst[1]
sheet['A3'] = lst[2]

workbook.save('example.xlsx')

# -------------------------------2--------------------------------

import json

with open("data.json", "r") as f:
    data = json.load(f)

for d in data:
    with open(f'{d["id"]}.json', 'w') as f:
        json.dump(d, f)


# -------------------------------3--------------------------------


import docx
from docx.shared import Pt 

dc=docx.Document("hello_python.docx")

for i in dc.paragraphs[0].runs:
    if i.bold:
        print(i.text)

doс = docx.Document()

text_to_add = "Это текст для нового абзаца в Word-документе."
doс.add_paragraph(text_to_add)

for run in doс.paragraphs[0].runs:
    run.font.size = Pt(22) 
    run.font.name = 'Bahnschrift'


doс.save("new.docx")
